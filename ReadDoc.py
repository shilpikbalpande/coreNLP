import docx
import os
import re
import textract
import win32com.client
from stanfordcorenlp import StanfordCoreNLP
regex_email = r'\w+[.|\w]\w+@\w+[.]\w+[.|\w+]\w+'
regex_ssn = r'(\d{3}-\d{2}-\d{4})'

file_path = r"C:\Users\i847808\Desktop"
filename = r"ShilpikBalpande_docx.docx"

class StanfordNLP:
  def __init__(self, host='http://localhost', port=9000):
    self.nlp = StanfordCoreNLP(host, port=port,
                               timeout=30000 )# , quiet =False, logging_level=logging.DEBUG)
    self.props = {
      'annotators': 'tokenize,ssplit,pos,lemma,ner,parse,depparse,dcoref,relation',
      'pipelineLanguage': 'en',
      'outputFormat': 'dict'
    }

  def ner(self, sentence):
    return self.nlp.ner(sentence)

if filename.endswith('.docx'):
  s = ''
  doc = docx.Document(file_path + "\\" + filename)
  sNLP = StanfordNLP()
  for i in doc.paragraphs:
    if i.text != "":
     s+= str(i.text) + "\t"

  for itemsInList in sNLP.ner(s):
   if itemsInList[1] == "PERSON":
    print (itemsInList[0])
   if bool(re.search(regex_email,s)) == True:   ## email
    print("Email address " + s)


elif filename.endswith('.doc'):
  app = win32com.client.Dispatch("Word.Application")
  app.Visible = False
  app.Documents.Open(file_path + "\\" + filename)
  doc = app.ActiveDocument
  docText = doc.Content.Text
  list = docText.splitlines()
  doc.Close()
  app.Quit()
     # converting .doc to .docx
  doc_file = file_path + "\\" + filename
  docx_file = file_path + "\\" + filename + 'x'
  if not os.path.exists(docx_file):
    os.system('antiword ' + doc_file + ' > ' + docx_file)
    text = textract.process(docx_file)
    for i in doc.paragraphs:
     print(i.text)
